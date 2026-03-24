import os
import sys
import re
import uuid
import json
import threading
import requests
import importlib
from datetime import datetime
from io import BytesIO
from flask import Flask, request, jsonify, send_from_directory

# --- CONFIGURAÇÕES DE AMBIENTE (Antigo run.py) ---
os.environ["PYTHONUNBUFFERED"] = "1" 
os.environ["PYTHONIOENCODING"] = "utf-8"

try:
    sys.stdout.reconfigure(encoding='utf-8')
    sys.stderr.reconfigure(encoding='utf-8')
except:
    pass

# --- DEPENDÊNCIAS DE TERCEIROS ---
# Certifique-se de ter instalado: pip install flask requests python-docx fpdf pillow
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from PIL import Image

# ==============================================================================
# 📋 CONFIGURAÇÕES E SITES (Antigo config.py)
# ==============================================================================

LOGO_PATH = r"C:\Users\Administrator\Desktop\datasheetget\utils\ventura.png"
# Altere para a pasta onde os arquivos finais devem cair:
OUTPUT_DIR = r"\\SERVIDOR2\Publico\Datasheet" 

HEADER_INFO = {
    "empresa": "VENTURA COMERCIO DE INFORMÁTICA EIRELI",
    "cnpj": "CNPJ: 08.310.365/0001-24",
    "endereco": "RUA SETE 560 COCAL VILA VELHA – ES I 29105-770"
}

URL_CALLBACK_PADRAO = "http://127.0.0.1:3000/api/datasheet/webhook"

SITES_CONFIG = {
    'MERCADO_LIVRE': {'padroes_url': [r'mercadolivre\.com', r'produto\.mercadolivre'], 'modulo': 'mercado_livre', 'classe': 'MercadoLivreScraper'},
    'AMAZON': {'padroes_url': [r'amazon\.com', r'amzn\.to'], 'modulo': 'amazon', 'classe': 'AmazonScraper'},
    'DELL': {'padroes_url': [r'dell\.com'], 'modulo': 'dell', 'classe': 'DellScraper'},
    'KABUM': {'padroes_url': [r'kabum\.com\.br'], 'modulo': 'kabum', 'classe': 'KabumScraper'},
    'BHPHOTOVIDEO': {'padroes_url': [r'bhphotovideo\.com'], 'modulo': 'bhphotovideo', 'classe': 'BhPhotoVideoScraper'},
    'FRIGELAR': {'padroes_url': [r'frigelar\.com\.br'], 'modulo': 'frigelar', 'classe': 'FrigelarScraper'},
    'samsung': {'padroes_url': [r'samsung\.com'], 'modulo': 'samsung', 'classe': 'SamsungScraper'}
    # Adicione os outros sites aqui conforme sua lista do config.py original
}

def identificar_site(url):
    for site_nome, config in SITES_CONFIG.items():
        for padrao in config['padroes_url']:
            if re.search(padrao, url, re.IGNORECASE):
                return site_nome, config['modulo'], config['classe']
    return None, None, None

# ==============================================================================
# 🎨 GERADOR DE DOCUMENTOS (Antigo generator.py)
# ==============================================================================

class DocGenerator:
    def _redimensionar_imagem(self, image_path):
        try:
            img = Image.open(image_path)
            if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                alpha = img.convert('RGBA').split()[-1]
                bg = Image.new("RGB", img.size, (255, 255, 255))
                bg.paste(img, mask=alpha)
                img = bg
            else:
                img = img.convert("RGB")
            MAX_SIZE = 500
            img.thumbnail((MAX_SIZE, MAX_SIZE), Image.Resampling.LANCZOS)
            square = Image.new('RGB', (MAX_SIZE, MAX_SIZE), (255, 255, 255))
            square.paste(img, ((MAX_SIZE - img.width) // 2, (MAX_SIZE - img.height) // 2))
            square.save(image_path, "JPEG", quality=95)
            return True
        except Exception as e:
            print(f"   ⚠️ Erro PIL: {e}", flush=True)
            return False

    def create_word(self, data, filepath):
        try:
            doc = Document()
            # Cabeçalho
            p = doc.add_paragraph()
            p.add_run(HEADER_INFO["empresa"] + "\n").bold = True
            p.add_run(f"{HEADER_INFO['cnpj']}\n{HEADER_INFO['endereco']}").font.size = Pt(9)
            
            doc.add_heading(data.get("titulo", "Produto"), 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            img_path = data.get("caminho_imagem_temp")
            if img_path and os.path.exists(img_path):
                self._redimensionar_imagem(img_path)
                doc.add_picture(img_path, width=Inches(3.2))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_heading('Descrição', level=1)
            doc.add_paragraph(data.get("descricao", "Sem descrição disponível."))
            
            specs = data.get("caracteristicas", {})
            if specs:
                doc.add_heading('Ficha Técnica', level=1)
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                for k, v in (specs.items() if isinstance(specs, dict) else specs):
                    row = table.add_row().cells
                    row[0].text = str(k)
                    row[1].text = str(v)
            doc.save(filepath)
            return True
        except Exception as e:
            print(f"   [ERRO WORD] {e}", flush=True)
            return False

    def create_pdf(self, data, filepath):
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            def txt(s): return str(s).encode('latin-1', 'replace').decode('latin-1')

            pdf.set_font("Helvetica", 'B', 10)
            pdf.cell(0, 5, txt(HEADER_INFO["empresa"]), ln=True)
            pdf.set_font("Helvetica", '', 8)
            pdf.cell(0, 4, txt(HEADER_INFO["cnpj"]), ln=True)
            pdf.cell(0, 4, txt(HEADER_INFO["endereco"]), ln=True)
            pdf.ln(10)

            pdf.set_font("Helvetica", 'B', 14)
            pdf.multi_cell(0, 8, txt(data.get("titulo", "Produto")), align='C')
            
            img_path = data.get("caminho_imagem_temp")
            if img_path and os.path.exists(img_path):
                pdf.image(img_path, x=65, w=80)
                pdf.ln(5)

            pdf.set_font("Helvetica", 'B', 11)
            pdf.cell(0, 8, txt("Descrição"), ln=True)
            pdf.set_font("Helvetica", '', 10)
            pdf.multi_cell(0, 5, txt(data.get("descricao", "")[:3000]))
            
            specs = data.get("caracteristicas", {})
            if specs:
                pdf.ln(5)
                pdf.set_font("Helvetica", 'B', 11)
                pdf.cell(0, 8, txt("Ficha Técnica"), ln=True)
                for k, v in (specs.items() if isinstance(specs, dict) else specs):
                    pdf.set_font("Helvetica", 'B', 9)
                    pdf.cell(0, 6, txt(str(k)), ln=True, fill=True)
                    pdf.set_font("Helvetica", '', 9)
                    pdf.multi_cell(0, 5, txt(str(v)), border='B')
            pdf.output(filepath)
            return True
        except Exception as e:
            print(f"   [ERRO PDF] {e}", flush=True)
            return False

# ==============================================================================
# 🚀 GERENCIADOR E API (Antigo scraper_manager.py e api.py)
# ==============================================================================

app = Flask(__name__)
pedidos = {}

# Garante que o Python veja a pasta 'scrapers' para importar o base e os outros
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

class ScraperManager:
    def executar_scraping(self, url, pasta_pedido):
        try:
            site_nome, modulo_nome, classe_nome = identificar_site(url)
            if not site_nome:
                return {'sucesso': False, 'erro': 'Site não configurado'}
            
            # Importa dinamicamente da pasta /scrapers
            modulo = importlib.import_module(f"scrapers.{modulo_nome}")
            ClasseScraper = getattr(modulo, classe_nome)
            
            scraper = ClasseScraper(url)
            scraper.output_folder = pasta_pedido
            print(f"   🔎 [ROBÔ] Iniciando: {site_nome}", flush=True)
            return scraper.executar()
        except Exception as e:
            return {'sucesso': False, 'erro': str(e)}

scraper_manager = ScraperManager()

def processar_pedido_background(request_id_interno, url, webhook_url, origem, codigo_tarefa, custom_id):
    try:
        id_bot = custom_id if custom_id else request_id_interno
        print(f" ⚙️ [THREAD] Processando ID: {id_bot}", flush=True)
        
        resultado = scraper_manager.executar_scraping(url, OUTPUT_DIR)
        
        base_url = "http://localhost:6004"
        files = resultado.get('arquivos', {})
        link_word = f"{base_url}/download/{files.get('word_nome')}" if files.get('word_nome') else None
        link_pdf = f"{base_url}/download/{files.get('pdf_nome')}" if files.get('pdf_nome') else None

        if resultado.get('sucesso'):
            print(f" ✅ [SUCESSO] Arquivos gerados para {id_bot}", flush=True)
            payload = {
                "status": "CONCLUIDO", "sucesso": True, "request_id": id_bot,
                "codigoTarefa": codigo_tarefa,
                "download": {"pdf": link_pdf, "word": link_word}
            }
        else:
            payload = {"status": "ERRO", "sucesso": False, "request_id": id_bot, "erro": resultado.get('erro')}

        pedidos[request_id_interno]['resultado'] = payload
        pedidos[request_id_interno]['status'] = 'concluido' if resultado.get('sucesso') else 'erro'

        if webhook_url:
            requests.post(webhook_url, json=payload, timeout=10)
    except Exception as e:
        print(f" 💥 [CRÍTICO] {e}", flush=True)

@app.route('/api/datasheet/processar', methods=['POST'])
def iniciar():
    print("\n >>> [REQUISIÇÃO RECEBIDA] <<<", flush=True)
    dados = request.get_json(silent=True, force=True)
    if not dados: return jsonify({"error": "JSON inválido"}), 400

    url = dados.get('url') or (dados.get('dados', {}).get('url') if 'dados' in dados else None)
    webhook = dados.get('webhook_url') or (dados.get('dados', {}).get('webhook_url') if 'dados' in dados else URL_CALLBACK_PADRAO)
    c_tarefa = dados.get('codigoTarefa')
    c_id = dados.get('custom_id') or dados.get('id')

    if not url: return jsonify({"error": "URL faltando"}), 400

    req_id = str(uuid.uuid4())
    pedidos[req_id] = {'status': 'processando'}

    thread = threading.Thread(target=processar_pedido_background, args=(req_id, url, webhook, "API", c_tarefa, c_id))
    thread.daemon = True
    thread.start()

    return jsonify({"success": True, "request_id": req_id}), 202

@app.route('/download/<path:filename>')
def download(filename):
    return send_from_directory(OUTPUT_DIR, filename, as_attachment=True)

@app.route('/health')
def health():
    return jsonify({"status": "online", "pasta": OUTPUT_DIR, "pedidos_ativos": len(pedidos)})

if __name__ == '__main__':
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print("="*60)
    print(f"🚀 DATASHEET MONOLITO ONLINE - PORTA 6004")
    print(f"📁 SALVANDO EM: {OUTPUT_DIR}")
    print("="*60)
    app.run(host='0.0.0.0', port=6004, threaded=True, debug=False)