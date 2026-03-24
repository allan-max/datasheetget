# utils/generator.py
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
from PIL import Image
import os

LOGO_PATH = r"C:\Users\Administrator\Desktop\datasheetget\utils\ventura.png"

HEADER_INFO = {
    "empresa": "VENTURA COMERCIO DE INFORMÁTICA EIRELI",
    "cnpj": "CNPJ: 08.310.365/0001-24",
    "endereco": "RUA SETE 560 COCAL VILA VELHA – ES I 29105-770"
}

class DocGenerator:
    def _redimensionar_imagem(self, image_path):
        print("   [FABRICA] Redimensionando a imagem para ficar quadrada...")
        try:
            img = Image.open(image_path)
            if img.mode in ('RGBA', 'LA') or (img.mode == 'P' and 'transparency' in img.info):
                alpha = img.convert('RGBA').split()[-1]
                bg = Image.new("RGB", img.size, (255, 255, 255))
                bg.paste(img, mask=alpha)
                img = bg
            else:
                img = img.convert("RGB")

            MAX_SIZE = 500
            img.thumbnail((MAX_SIZE, MAX_SIZE), Image.Resampling.LANCZOS)
            square = Image.new('RGB', (MAX_SIZE, MAX_SIZE), (255, 255, 255))
            offset_x = (MAX_SIZE - img.width) // 2
            offset_y = (MAX_SIZE - img.height) // 2
            square.paste(img, (offset_x, offset_y))
            square.save(image_path, "JPEG", quality=95)
            print("   [FABRICA] Imagem redimensionada com sucesso!")
            return True
        except Exception as e:
            print(f"   [ERRO FABRICA] Falha ao tratar a imagem: {e}")
            return False

    def _adicionar_cabecalho_word(self, doc):
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False
        table.columns[0].width = Inches(4.5) 
        table.columns[1].width = Inches(2.0)
        cell_text = table.cell(0, 0)
        p = cell_text.paragraphs[0]
        run_nome = p.add_run(HEADER_INFO["empresa"] + "\n")
        run_nome.bold = True
        run_nome.font.size = Pt(11)
        run_nome.font.name = 'Arial'
        run_rest = p.add_run(f"{HEADER_INFO['cnpj']}\n{HEADER_INFO['endereco']}")
        run_rest.font.size = Pt(9)
        run_rest.font.name = 'Arial'
        cell_img = table.cell(0, 1)
        p_img = cell_img.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if os.path.exists(LOGO_PATH):
            try: run_img = p_img.add_run().add_picture(LOGO_PATH, width=Inches(0.9))
            except: pass
        doc.add_paragraph("")

    def create_word(self, data, filepath):
        print(f"   [FABRICA] Escrevendo documento Word em: {os.path.basename(filepath)}")
        try:
            doc = Document()
            self._adicionar_cabecalho_word(doc)
            h = doc.add_heading(data.get("titulo", "Produto"), 0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER

            img_path = data.get("caminho_imagem_temp")
            if img_path and os.path.exists(img_path):
                self._redimensionar_imagem(img_path)
                try:
                    doc.add_picture(img_path, width=Inches(3.2))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                except: pass
            
            doc.add_heading('Descrição', level=1)
            doc.add_paragraph(data.get("descricao", "Sem descrição disponível."))

            specs = data.get("caracteristicas", {})
            if specs:
                doc.add_heading('Ficha Técnica', level=1)
                table = doc.add_table(rows=0, cols=2)
                table.style = 'Table Grid'
                hdr = table.add_row().cells
                hdr[0].text = 'Item'
                hdr[1].text = 'Detalhe'
                items = specs.items() if isinstance(specs, dict) else specs
                for k, v in items:
                    row = table.add_row().cells
                    row[0].text = str(k)
                    row[1].text = str(v)

            doc.save(filepath)
            print("   [FABRICA] Word salvo com sucesso!")
            return True
        except Exception as e:
            print(f"   [ERRO FABRICA] Falha no Word: {e}")
            return False

    def create_pdf(self, data, filepath):
        print(f"   [FABRICA] Desenhando PDF em: {os.path.basename(filepath)}")
        try:
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            def txt(s):
                s = str(s).replace('’', "'").replace('“', '"').replace('”', '"')
                return s.encode('latin-1', 'replace').decode('latin-1')

            if os.path.exists(LOGO_PATH):
                try: pdf.image(LOGO_PATH, x=165, y=8, w=30)
                except: pass
            pdf.set_xy(10, 10)
            pdf.set_font("Helvetica", 'B', 10)
            pdf.cell(0, 5, txt(HEADER_INFO["empresa"]), ln=True)
            pdf.set_font("Helvetica", '', 8)
            pdf.cell(0, 4, txt(HEADER_INFO["cnpj"]), ln=True)
            pdf.cell(0, 4, txt(HEADER_INFO["endereco"]), ln=True)
            pdf.ln(8)

            pdf.set_font("Helvetica", 'B', 14)
            pdf.multi_cell(0, 8, txt(data.get("titulo", "Produto")), align='C')
            pdf.ln(5)

            img_path = data.get("caminho_imagem_temp")
            if img_path and os.path.exists(img_path):
                try:
                    x_pos = (210 - 80) / 2
                    pdf.image(img_path, x=x_pos, w=80)
                    pdf.ln(5)
                except: pass

            pdf.set_font("Helvetica", 'B', 11)
            pdf.cell(0, 8, txt("Descrição"), ln=True)
            pdf.set_font("Helvetica", '', 10)
            pdf.multi_cell(0, 5, txt(data.get("descricao", "")[:3000]))
            pdf.ln(5)

            specs = data.get("caracteristicas", {})
            if specs:
                pdf.set_font("Helvetica", 'B', 11)
                pdf.cell(0, 8, txt("Ficha Técnica"), ln=True)
                items = specs.items() if isinstance(specs, dict) else specs
                for k, v in items:
                    pdf.set_font("Helvetica", 'B', 9)
                    pdf.set_fill_color(240, 240, 240)
                    pdf.cell(0, 6, txt(str(k)), ln=True, fill=True)
                    pdf.set_font("Helvetica", '', 9)
                    pdf.multi_cell(0, 5, txt(str(v)), border='B')
                    pdf.ln(1)

            pdf.output(filepath)
            print("   [FABRICA] PDF gerado com sucesso!")
            return True
        except Exception as e:
            print(f"   [ERRO FABRICA] Falha no PDF: {e}")
            return False